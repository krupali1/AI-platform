"""
All pandas/openpyxl file reading lives here - isolates the one heavy
dependency this feature adds, and keeps tally_pipeline.py free of
file-format concerns.
"""
import io
import re
import pandas as pd


def parse_excel_file(blob, content_type=None, filename=None):
    """Returns {sheet_name: DataFrame} for an uploaded .xlsx/.xls/.csv
    blob. A CSV has no sheet name, so it's given the synthetic name
    "Sheet1" to keep every caller dealing with the same shape."""
    buf = io.BytesIO(blob)
    name = (filename or "").lower()
    if name.endswith(".csv") or (content_type and "csv" in content_type):
        df = pd.read_csv(buf, dtype=str, keep_default_na=False)
        return {"Sheet1": df}
    sheets = pd.read_excel(buf, sheet_name=None, dtype=str, keep_default_na=False)
    return sheets


def sheet_names_of(parsed):
    return list(parsed.keys())


def column_names_of(parsed, sheet_name):
    if sheet_name not in parsed:
        return []
    return list(parsed[sheet_name].columns)


def build_master_sku_map(parsed, sheet_name, sku_column, code_column, multiplier_column=None):
    """Platform SKU/ASIN -> a list of {code, multiplier} components, from
    the parsed Master file. A list rather than a single code is what
    makes combo/bundle packs pure master-file configuration: a same-
    product combo (SKU "X_PK3" = 3x the base product) is one master row
    with multiplier 3; a combo of different products is multiple master
    rows sharing the same SKU, one per component - PIL adds a row to
    the master sheet for a new combo, never touches code or a rule.
    Blank/missing SKUs are skipped rather than mapped to an empty
    string, so a lookup miss stays a genuine miss.

    multiplier is None (not defaulted to 1) when the column isn't
    mapped or a row's value is blank/invalid - resolving a blank into
    an actual number is tally_pipeline.py's job (master-file value if
    present, else a SKU-suffix pattern, else 1), not this parsing
    layer's - confirmed against real PIL data that the Master file's
    own multiplier column is usually left blank even for genuine
    combo SKUs, so "blank" and "explicitly 1" have to stay
    distinguishable this far down.

    A SKU whose (code, multiplier) pair repeats identically across more
    than one master row is NOT a combo of identical components - it's
    the same fact stated twice, a real data-entry duplication confirmed
    in a real PIL master file (several SKUs each listed twice with the
    exact same Product Name and a blank Pack Of). Left undeduped, this
    silently doubled the affected orders' output rows, since the combo-
    unbundling logic downstream has no way to tell "genuinely 2
    identical units" apart from "the same row entered twice." A real
    same-product combo is still expressed correctly - as a single row
    with an explicit multiplier, or via the SKU-suffix pattern - so
    deduping exact repeats here never discards a real multi-unit pack."""
    if sheet_name not in parsed:
        return {}
    df = parsed[sheet_name]
    if sku_column not in df.columns or code_column not in df.columns:
        return {}
    has_multiplier_col = multiplier_column and multiplier_column in df.columns
    mapping = {}
    for _, row in df.iterrows():
        sku = str(row[sku_column]).strip()
        code = str(row[code_column]).strip()
        if not sku or not code:
            continue
        multiplier = None
        if has_multiplier_col:
            raw = str(row[multiplier_column]).strip()
            if raw:
                try:
                    multiplier = float(raw)
                except ValueError:
                    multiplier = None
        component = {"code": code, "multiplier": multiplier}
        components = mapping.setdefault(sku, [])
        if component not in components:
            components.append(component)
    return mapping


def build_batch_queue(parsed, sheet_name, product_column, location_column, batch_column, qty_column, mfg_column, exp_column):
    """(internal_product_code, location) -> [{batch_no, qty, mfg_date,
    exp_date}, ...] from the parsed Batch wise Summary file, sorted
    oldest MFG date first - the PRD's own FIFO example ("Batch 1 was
    created before Batch 2") ties "created before" to manufacture date,
    since the file has no separate creation-date column. A row with a
    blank product/batch/qty is skipped rather than guessed into a
    queue entry; a blank MFG date sorts last (treated as "unknown
    recency", never assumed oldest) rather than crashing the sort or
    silently being treated as the newest/oldest batch.

    location_column is optional - the PRD's own description of this
    logic ("the stock is picked in a FIFO manner") never scopes it by
    location, only by product, so a Batch wise Summary file with no
    location/warehouse column still allocates correctly: every batch
    for a product pools into one queue keyed by (product, ""). When a
    location column IS mapped, matching is scoped to (product,
    location) instead, for a file that does track batches per
    warehouse."""
    if sheet_name not in parsed:
        return {}
    df = parsed[sheet_name]
    required = (product_column, batch_column, qty_column)
    if any(c not in df.columns for c in required):
        return {}
    has_location = location_column and location_column in df.columns
    has_mfg = mfg_column and mfg_column in df.columns
    has_exp = exp_column and exp_column in df.columns
    queues = {}
    for _, row in df.iterrows():
        product = str(row[product_column]).strip()
        location = str(row[location_column]).strip() if has_location else ""
        batch_no = str(row[batch_column]).strip()
        if not product or not batch_no or (has_location and not location):
            continue
        try:
            qty = float(str(row[qty_column]).strip())
        except (TypeError, ValueError):
            continue
        mfg_date = str(row[mfg_column]).strip() if has_mfg else ""
        exp_date = str(row[exp_column]).strip() if has_exp else ""
        queues.setdefault((product, location), []).append({
            "batch_no": batch_no, "qty": qty, "mfg_date": mfg_date, "exp_date": exp_date,
        })
    for key, batches in queues.items():
        batches.sort(key=lambda b: (b["mfg_date"] == "", b["mfg_date"]))
    return queues


_STOCK_SUMMARY_HEADER_TEXT = "Opening Balance"
_MFG_EXP_PATTERN = re.compile(r"mfg\s*date\s*:\s*([^e]*?)\s*expiry\s*date\s*:\s*(.*)", re.IGNORECASE)
# PIL's own confirmed alias: the "Main Location" Stock Summary export's
# title has no warehouse code Amazon's Warehouse Id would ever contain -
# it's PIL's own godown (Amazon Seller Flex), which the rest of the
# pipeline already knows by the code "VZPL".
_STOCK_SUMMARY_LOCATION_ALIASES = {"MAIN": "VZPL"}


def parse_tally_stock_summary(blob):
    """Parses a native Tally "Stock Summary" report export - the real
    shape PIL's Batch wise Summary files turned out to be: one file per
    warehouse, exported straight out of Tally (File > Export), not a
    spreadsheet PIL typed by hand. Returns None if the file doesn't
    look like this shape at all (no "Opening Balance" group header
    found in the first few rows of any sheet), so the caller can fall
    back to treating it as a flat, user-mapped sheet instead - this
    format needs no column mapping at all, since its layout is fixed
    and self-describing.

    Layout: a title block whose first line is "<company name> <location
    code>" (e.g. "PSYCHOTROPICS INDIA LIMITED BLR-7" -> location code
    "BLR-7" - the last whitespace-separated token), then a two-row
    grouped header (Opening Balance / Inwards / Outwards / Closing
    Balance, each with its own Quantity/Alt. Units/Rate/Value sub-
    columns), then one row per stock item (Excel cell indent 0)
    followed by one indented row per batch of that item (indent > 0).
    MFG/EXP dates are only present on a batch row that tracks them,
    packed into one text cell as "Mfg Date :X  Expiry Date :Y". Data
    ends at the "Grand Total" row.

    Only the Opening Balance group's Quantity column is read - PIL
    confirmed this is the stock actually available to allocate this
    period's sales against; Inwards/Outwards/Closing reflect movements
    Tally has already recorded elsewhere, not stock this pipeline's own
    FIFO allocation should start consuming from scratch."""
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(blob), data_only=True)
    for ws in wb.worksheets:
        result = _parse_stock_summary_sheet(ws)
        if result is not None:
            return result
    return None


def _parse_stock_summary_sheet(ws):
    max_row, max_col = ws.max_row, ws.max_column
    header_row = None
    for r in range(1, min(max_row, 30) + 1):
        for c in range(1, max_col + 1):
            v = ws.cell(row=r, column=c).value
            if isinstance(v, str) and v.strip() == _STOCK_SUMMARY_HEADER_TEXT:
                header_row = r
                break
        if header_row:
            break
    if not header_row:
        return None

    # Forward-fill the top header row's group names across the blank
    # cells beneath a merged/spanning group label, so every column can
    # be attributed to its group (Opening Balance / Inwards / ...).
    group_by_col = {}
    current_group = None
    for c in range(1, max_col + 1):
        v = ws.cell(row=header_row, column=c).value
        if isinstance(v, str) and v.strip():
            current_group = v.strip()
        group_by_col[c] = current_group

    qty_col = None
    for c in range(1, max_col + 1):
        sub = ws.cell(row=header_row + 1, column=c).value
        if group_by_col.get(c) == _STOCK_SUMMARY_HEADER_TEXT and isinstance(sub, str) and sub.strip() == "Quantity":
            qty_col = c
            break
    if not qty_col:
        return None

    location_code = None
    for r in range(1, header_row):
        v = ws.cell(row=r, column=1).value
        if isinstance(v, str) and v.strip():
            tokens = v.strip().split()
            if tokens:
                location_code = tokens[-1]
            break
    if location_code:
        normalized = re.sub(r"[^A-Za-z0-9]", "", location_code).upper()
        location_code = _STOCK_SUMMARY_LOCATION_ALIASES.get(normalized, normalized)

    batches_by_product = {}
    current_product = None
    for r in range(header_row + 2, max_row + 1):
        cell0 = ws.cell(row=r, column=1)
        label = cell0.value
        if label is None:
            continue
        label = str(label).strip()
        if not label:
            continue
        if label.lower() == "grand total":
            break
        indent = (cell0.alignment.indent if cell0.alignment else 0) or 0
        if indent == 0:
            current_product = label
            continue
        if not current_product:
            continue
        qty_raw = ws.cell(row=r, column=qty_col).value
        try:
            qty = float(qty_raw) if qty_raw not in (None, "") else 0.0
        except (TypeError, ValueError):
            qty = 0.0
        if qty <= 0:
            continue
        mfg_date, exp_date = "", ""
        detail = ws.cell(row=r, column=2).value
        if isinstance(detail, str) and detail.strip():
            m = _MFG_EXP_PATTERN.search(detail)
            if m:
                mfg_date, exp_date = m.group(1).strip(), m.group(2).strip()
        batches_by_product.setdefault(current_product, []).append({
            "batch_no": label, "qty": qty, "mfg_date": mfg_date, "exp_date": exp_date,
        })
    return {"location_code": location_code, "batches_by_product": batches_by_product}


def rows_of(parsed, sheet_name):
    """Yields each row of a sheet as a plain dict, in file order, with
    a 1-based row index for traceability back to the raw upload."""
    if sheet_name not in parsed:
        return
    df = parsed[sheet_name]
    for idx, row in df.iterrows():
        yield idx + 1, row.to_dict()


DOCUMENT_EXTENSIONS = (".txt", ".md", ".docx", ".pdf", ".xlsx", ".xls", ".csv")


def extract_document_text(blob, content_type=None, filename=None):
    """Plain text out of a rules/conditions document a human wrote in
    whatever format they already had it in - a Word doc, a PDF, a
    plain text/markdown note, or a spreadsheet of edge cases. Used only
    to feed tally_rules.suggest_rules_from_document(); never parsed as
    structured data the way an Amazon/Master/sample file is, so this
    lives here (the file-format module) rather than in tally_rules.py
    (which stays free of format concerns, same reasoning as
    parse_excel_file living here instead of in tally_pipeline.py)."""
    name = (filename or "").lower()

    if name.endswith(".docx"):
        import docx
        document = docx.Document(io.BytesIO(blob))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    if name.endswith(".pdf"):
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(blob))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if name.endswith((".xlsx", ".xls", ".csv")):
        parsed = parse_excel_file(blob, content_type, filename)
        lines = []
        for sheet in sheet_names_of(parsed):
            lines.append(f"# {sheet}")
            for _, row in rows_of(parsed, sheet):
                cells = [f"{k}: {v}" for k, v in row.items() if str(v).strip()]
                if cells:
                    lines.append(", ".join(cells))
        return "\n".join(lines)

    # .txt / .md / anything else - treat as plain text
    return blob.decode("utf-8", errors="replace")
